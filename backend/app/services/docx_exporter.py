from docx import Document
from app.schemas.question import QuestionPreview
import io

def generate_docx_summary(questions: list, resumo: str, logo_path: str = None) -> bytes:
    doc = Document()
    if logo_path:
        doc.sections[0].header.add_paragraph().add_run().add_picture(logo_path, width=2)
    doc.add_heading("Resumo Inteligente", level=1)
    doc.add_paragraph(resumo)
    doc.add_heading("Questões", level=2)
    for q in questions:
        doc.add_paragraph(f"{q.numero}) {q.enunciado}", style="List Number")
    doc.add_page_break()
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
